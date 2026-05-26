from docx import Document
from docx.shared import Pt

# --- RAPORT TEMPLATE ---
doc_raport = Document(r"C:\Users\evn\Videos\aplikasi raport tk pgri\isi raport\1. BILKIS ANAJWA.docx")

replacements_raport = {
    "BILKIS ANAJWA": "{nama_siswa}",
    "0636": "{nipd}",
    "101": "{tinggi_badan}",
    "14,20": "{berat_badan}",
    "I / 2025-2026": "{semester}",
    ": 8": ": {sakit}",
    ": 1": ": {izin}",
    ": -": ": {tanpa_keterangan}",
    "Tuban, 20 Desember 2025": "{tanggal_raport}",
    "NamaSekolah": "{nama_sekolah}",
    "NamaSiswa": "Nama Siswa",
    "TUNIK LUSTARI, S.Pd": "{guru_kelas}",
    "NPA. 13101200817": "NPA. {npa_guru}",
    "INDAH ROHMAWATI, S.Pd": "{kepala_tk}",
    "NPA. 13101200816": "NPA. {npa_kepala}",
    "YAYASAN PEMBINA LEMBAGA PENDIDIKAN": "{kop_1}",
    "PERSATUAN GURU REPUBLIK INDONESIA JAWA TIMUR": "{kop_2}",
    "(YPLP PGRI JATIM) PERWAKILAN KABUPATEN TUBAN": "{kop_3}",
    "TAMAN KANAK-KANAK PGRI NUR IKHLAS": "{kop_4}",
    "DESA PRUNGGAHAN KULON KECAMATAN SEMANDING": "{kop_5}",
    "NPSN : 20574036 Email: tkpgrinuriklas@gmail.com": "{kop_6}"
}

ai_texts = {
    "Alhamdulillah di semester ini ananda BILKIS": "{teks_agama}",
    "Ananda untuk perkembangan jati diri misalnya": "{teks_jati_diri}",
    "Ananda dalam perkembangan literasi atau bahasa": "{teks_literasi}",
    "Semester ini Ananda melakukan projek": "{teks_projek}"
}

loop_tags = ["loop_agama", "loop_jati_diri", "loop_literasi", "loop_projek"]
foto_idx = 0

def replace_preserve_format(para, search, replace, default_font="Calibri", default_size=14):
    if search in para.text:
        # Calculate new text
        new_text = para.text.replace(search, replace)
        
        # Safely modify text elements without destroying non-text elements (like <w:drawing> / logos)
        t_elements = para._element.xpath('.//w:t')
        if t_elements:
            t_elements[0].text = new_text
            for t in t_elements[1:]:
                t.text = ''
        
        # Enforce formatting on all runs
        for run in para.runs:
            # Only force font on runs that actually have text
            # (We don't want to mess with drawing runs if not needed, though it's usually harmless)
            run.font.name = default_font
            
            if "{semester}" in para.text:
                run.font.size = Pt(12)
            else:
                run.font.size = Pt(default_size)

for para in doc_raport.paragraphs:
    for search, replace in replacements_raport.items():
        replace_preserve_format(para, search, replace, default_font="Calibri", default_size=14)
    for search, replace in ai_texts.items():
        replace_preserve_format(para, search, replace, default_font="Calibri", default_size=14)

for table in doc_raport.tables:
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            # Process text replacements
            for para in cell.paragraphs:
                for search, replace in replacements_raport.items():
                    replace_preserve_format(para, search, replace, default_font="Calibri", default_size=14)
                
                for search, replace in ai_texts.items():
                    replace_preserve_format(para, search, replace, default_font="Calibri", default_size=14)
            
            # Process images
            cell_text = cell.text.strip().upper()
            if cell_text == "FOTO KEGIATAN" or cell_text == "FOTO KEGIATAN ANAK":
                if r_idx + 1 < len(table.rows):
                    target_cell = table.rows[r_idx + 1].cells[c_idx]
                    if foto_idx < len(loop_tags):
                        loop_tag = loop_tags[foto_idx]
                        
                        target_cell.text = "" # Clear old photos
                        para = target_cell.paragraphs[0]
                        para.alignment = 1 # Center alignment
                        
                        # Add tags with simple horizontal space
                        para.add_run(f"{{#{loop_tag}}}")
                        para.add_run("  ") 
                        para.add_run("{%foto_img}")
                        para.add_run("  ") 
                        para.add_run(f"{{/{loop_tag}}}")
                        
                        foto_idx += 1

doc_raport.save("public/template_raport_v2.docx")
print("Template Raport saved.")


# --- IDENTITAS TEMPLATE ---
doc_depan = Document(r"C:\Users\evn\Videos\aplikasi raport tk pgri\LEMBAR DEPAN SAMA LEMBAR IDENTITAS.docx")

replacements_depan = {
    "MUHAMMAD ABDULLAH": "{nama_lengkap}",
    "AHMAD": "{nama_panggilan}",
    "0706 / 3202211551": "{nisn}",
    "Laki-laki": "{jenis_kelamin}",
    "Tuban, 16 Mei 2020": "{tempat_tanggal_lahir}",
    "Islam": "{agama}",
    "1 (Satu)": "{anak_ke}",
    "GISO SISWOKO": "{nama_ayah}",
    "ROSITA SARI SIREGAR": "{nama_ibu}",
    "Wiraswasta": "{pekerjaan_ayah}",
    "Mengurus Rumah Tangga": "{pekerjaan_ibu}",
    "Dsn. Mojokopek  Rt.002 / Rw.030": "{alamat_jalan}",
    "081359092538": "{telepon}",
    "Semanding, 14 Juli 2025": "{tanggal_identitas}",
    "Prunggahan Kulon": "{desa}",
    "Semanding": "{kecamatan}",
    "Tuban": "{kabupaten}",
    "Jawa Timur": "{provinsi}",
    
    # Identitas Lembaga
    "PGRI NUR IKHLAS": "{nama_tk_lembaga}",
    "004050603035 / 20574036": "{nss_npsn_lembaga}",
    "DSN. MOJOKOPEK RT.01/RW.29": "{alamat_tk_lembaga}",
    "62381": "{kode_pos_lembaga}",
    "PRUNGGAHAN KULON": "{desa_lembaga}",
    "SEMANDING": "{kec_lembaga}",
    "TUBAN": "{kab_lembaga}",
    "JAWA TIMUR": "{prov_lembaga}",

    "NamaSekolah": "{nama_sekolah}",
    "NamaSiswa": "Nama Siswa",
    "YAYASAN PEMBINA LEMBAGA PENDIDIKAN": "{kop_1}",
    "PERSATUAN GURU REPUBLIK INDONESIA JAWA TIMUR": "{kop_2}",
    "(YPLP PGRI JATIM) PERWAKILAN KABUPATEN TUBAN": "{kop_3}",
    "TAMAN KANAK-KANAK PGRI NUR IKHLAS": "{kop_4}",
    "DESA PRUNGGAHAN KULON KECAMATAN SEMANDING": "{kop_5}",
    "NPSN : 20574036 Email: tkpgrinuriklas@gmail.com": "{kop_6}"
}

for para in doc_depan.paragraphs:
    for search, replace in replacements_depan.items():
        replace_preserve_format(para, search, replace, default_font="Arial", default_size=14)

for table in doc_depan.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for search, replace in replacements_depan.items():
                    replace_preserve_format(para, search, replace, default_font="Arial", default_size=14)

doc_depan.save("public/template_depan_v2.docx")
print("Template Depan saved.")
