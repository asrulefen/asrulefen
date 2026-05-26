from docx import Document
import re

doc_path = r"C:\Users\evn\Videos\aplikasi raport tk pgri\app-raport\public\template_raport_v2.docx"
doc = Document(doc_path)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "{%foto_img}" in para.text:
                    # Found a corrupted photo paragraph!
                    text = para.text
                    match = re.search(r'\{#(loop_[a-z_]+)\}', text)
                    if match:
                        loop_tag = match.group(1)
                        
                        # Save paragraph style/alignment if needed
                        alignment = para.alignment
                        
                        # Clear all corrupted merged runs
                        para.clear()
                        
                        if alignment is not None:
                            para.alignment = alignment
                        else:
                            para.alignment = 1 # Center
                        
                        # Rebuild cleanly isolated runs with NO spaces. 
                        # We will handle spacing via image padding in the frontend canvas!
                        para.add_run(f"{{#{loop_tag}}}")
                        para.add_run("{%foto_img}")
                        para.add_run(f"{{/{loop_tag}}}")

doc.save(doc_path)
print("Fixed manual edits XML corruption!")
