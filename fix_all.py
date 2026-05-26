from docx import Document
from docx.shared import Pt

# --- Fix Template Raport (Isi Raport) ---
doc_raport = Document(r"C:\Users\evn\Videos\aplikasi raport tk pgri\isi raport\1. BILKIS ANAJWA.docx")

replacements_raport = {
    "101": "{tinggi_badan}",
    "14,20": "{berat_badan}",
    "Sakit: 8": "Sakit: {sakit}",
    "Izin: 1": "Izin: {izin}",
    "Tanpa Keterangan: -": "Tanpa Keterangan: {tanpa_keterangan}",
    "NamaSekolah": "Nama Sekolah",
    "NamaSiswa": "Nama Siswa",
}

ai_texts = {
    "Alhamdulillah di semester ini ananda BILKIS": "{teks_agama}",
    "Ananda untuk perkembangan jati diri misalnya": "{teks_jati_diri}",
    "Ananda dalam perkembangan literasi atau bahasa": "{teks_literasi}",
    "Semester ini Ananda melakukan projek": "{teks_projek}"
}

loop_tags = ["loop_agama", "loop_jati_diri", "loop_literasi", "loop_projek"]
foto_idx = 0

def apply_font(para):
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)

def process_paragraphs_raport(paragraphs):
    for para in paragraphs:
        for search, replace in replacements_raport.items():
            if search in para.text:
                para.text = para.text.replace(search, replace)
                
        for search, replace in ai_texts.items():
            if search in para.text:
                para.text = replace
        apply_font(para)

process_paragraphs_raport(doc_raport.paragraphs)

for table in doc_raport.tables:
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            process_paragraphs_raport(cell.paragraphs)
            
            # Check if this cell is the "FOTO KEGIATAN" header
            cell_text = cell.text.strip().upper()
            if cell_text == "FOTO KEGIATAN" or cell_text == "FOTO KEGIATAN ANAK":
                if r_idx + 1 < len(table.rows):
                    target_cell = table.rows[r_idx + 1].cells[c_idx]
                    if foto_idx < len(loop_tags):
                        loop_tag = loop_tags[foto_idx]
                        target_cell.text = ""
                        para = target_cell.paragraphs[0]
                        para.alignment = 1 # Center
                        
                        # Add extra line spacing for vertical gaps between photos!
                        para.paragraph_format.line_spacing = 1.2
                        
                        run1 = para.add_run(f"{{#{loop_tag}}}")
                        run2 = para.add_run("  ") # Horizontal spacing
                        run3 = para.add_run("{%foto_img}")
                        run4 = para.add_run("  ") # Horizontal spacing
                        run5 = para.add_run(f"{{/{loop_tag}}}")
                        
                        apply_font(para)
                        foto_idx += 1

doc_raport.save("public/template_raport_v2.docx")
print("Template Raport saved.")

# --- Fix Template Depan (Identitas) ---
doc_depan = Document(r"C:\Users\evn\Videos\aplikasi raport tk pgri\LEMBAR DEPAN SAMA LEMBAR IDENTITAS.docx")

replacements_depan = {
    "MUHAMMAD ABDULLAH": "{nama_lengkap}",
    "AHMAD": "{nama_panggilan}",
    "0706 / 3202211551": "{nisn}",
    "Laki-laki": "{jenis_kelamin}",
    "Tuban, 16 Mei 2020": "{tempat_tanggal_lahir}",
    "Islam": "{agama}",
    "1 (Satu)": "{anak_ke}",
    "GISO SISWOKO": "{nama_ayah}",
    "ROSITA SARI SIREGAR": "{nama_ibu}",
    "Wiraswasta": "{pekerjaan_ayah}",
    "Mengurus Rumah Tangga": "{pekerjaan_ibu}",
    "Dsn. Mojokopek  Rt.002 / Rw.030": "{alamat_jalan}",
    "081359092538": "{telepon}",
    "Prunggahan Kulon": "{desa}",
    "Semanding": "{kecamatan}",
    "Tuban": "{kabupaten}",
    "Jawa Timur": "{provinsi}",
    "NamaSekolah": "Nama Sekolah",
    "NamaSiswa": "Nama Siswa",
}

def process_paragraphs_depan(paragraphs):
    for para in paragraphs:
        for search, replace in replacements_depan.items():
            if search in para.text:
                para.text = para.text.replace(search, replace)
        apply_font(para)

process_paragraphs_depan(doc_depan.paragraphs)

for table in doc_depan.tables:
    for row in table.rows:
        for cell in row.cells:
            process_paragraphs_depan(cell.paragraphs)

doc_depan.save("public/template_depan_v2.docx")
print("Template Depan saved.")
