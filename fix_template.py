from docx import Document

doc = Document(r"C:\Users\evn\Videos\aplikasi raport tk pgri\isi raport\1. BILKIS ANAJWA.docx")

replacements = {
    "101": "{tinggi_badan}",
    "14,20": "{berat_badan}",
    "Sakit: 8": "Sakit: {sakit}",
    "Izin: 1": "Izin: {izin}",
    "Tanpa Keterangan: -": "Tanpa Keterangan: {tanpa_keterangan}",
}

ai_texts = {
    "Alhamdulillah di semester ini ananda BILKIS": "{teks_agama}",
    "Ananda untuk perkembangan jati diri misalnya": "{teks_jati_diri}",
    "Ananda dalam perkembangan literasi atau bahasa": "{teks_literasi}",
    "Semester ini Ananda melakukan projek": "{teks_projek}"
}

loop_tags = ["loop_agama", "loop_jati_diri", "loop_literasi", "loop_projek"]
foto_idx = 0

def process_paragraphs(paragraphs):
    for para in paragraphs:
        for search, replace in replacements.items():
            if search in para.text:
                para.text = para.text.replace(search, replace)
                
        for search, replace in ai_texts.items():
            if search in para.text:
                para.text = replace

process_paragraphs(doc.paragraphs)

for table in doc.tables:
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            process_paragraphs(cell.paragraphs)
            
            # Check if this cell is the "FOTO KEGIATAN" header
            cell_text = cell.text.strip().upper()
            if cell_text == "FOTO KEGIATAN" or cell_text == "FOTO KEGIATAN ANAK":
                # The photos are in the cell directly below this one
                if r_idx + 1 < len(table.rows):
                    target_cell = table.rows[r_idx + 1].cells[c_idx]
                    
                    if foto_idx < len(loop_tags):
                        loop_tag = loop_tags[foto_idx]
                        
                        # Clear existing photos
                        target_cell.text = ""
                        
                        # Inject new tags in completely separated runs
                        para = target_cell.paragraphs[0]
                        para.alignment = 1 # Center alignment (WD_ALIGN_PARAGRAPH.CENTER)
                        para.add_run(f"{{#{loop_tag}}}")
                        para.add_run(" ") # Spacing between photos if they end up inline
                        para.add_run("{%foto_img}")
                        para.add_run(" ") # Spacing
                        para.add_run(f"{{/{loop_tag}}}")
                        
                        foto_idx += 1

doc.save("public/template_raport_v2.docx")
print("Template saved to public/template_raport_v2.docx")
