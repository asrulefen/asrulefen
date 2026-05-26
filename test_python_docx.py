from docx import Document

doc = Document(r"C:\Users\evn\Videos\aplikasi raport tk pgri\isi raport\1. BILKIS ANAJWA.docx")
found = 0
for para in doc.paragraphs:
    if "Alhamdulillah" in para.text:
        print("Found Alhamdulillah in paragraph:", para.text[:50])
        found += 1
print("Total found:", found)

# Let's also check tables!
print("\nChecking tables...")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "Alhamdulillah" in para.text:
                    print("Found in table:", para.text[:50])
                    found += 1

print("Total found overall:", found)
